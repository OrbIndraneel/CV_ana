"""
docx_parser.py
DOCX parser module using python-docx.
Extracts standard body text, headers/footers, and flags internal formatting risks.
"""

import io
import logging
from typing import Any, Dict
import docx

from app.core.exceptions import NLPProcessingException
from app.nlp.utils.text_cleaner import clean_text

logger = logging.getLogger(__name__)


def parse_docx(file_bytes: bytes) -> Dict[str, Any]:
    """
    Parse text from a DOCX binary stream.
    Flags whether tables exist in the document, and scans header/footer blocks.
    """
    if not file_bytes:
        raise NLPProcessingException("Empty file bytes provided for DOCX parsing.")

    raw_text_parts = []
    tables_detected = False
    header_footer_text_parts = []

    try:
        # Load document in-memory
        doc = docx.Document(io.BytesIO(file_bytes))

        # 1. Parse main paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                raw_text_parts.append(paragraph.text)

        # 2. Check for tables
        if len(doc.tables) > 0:
            tables_detected = True
            logger.debug(f"Detected {len(doc.tables)} tables in DOCX document.")
            # We extract text inside table cells to ensure we don't lose any content.
            # However, we flag the tables as an ATS risk.
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text:
                            # Strip to prevent formatting spillover
                            row_text.append(cell.text.strip())
                    # Join cell text with spacing
                    if row_text:
                        raw_text_parts.append(" | ".join(row_text))

        # 3. Parse headers and footers for contact details (often skipped by simple parsers)
        for section in doc.sections:
            # Check header
            if section.header and not section.header.is_linked_to_previous:
                for p in section.header.paragraphs:
                    if p.text and p.text.strip():
                        header_footer_text_parts.append(p.text)
            # Check footer
            if section.footer and not section.footer.is_linked_to_previous:
                for p in section.footer.paragraphs:
                    if p.text and p.text.strip():
                        header_footer_text_parts.append(p.text)

        # Combine document blocks
        full_text = "\n".join(raw_text_parts)

        # If headers or footers had content, append it with warning markers
        header_footer_text = "\n".join(header_footer_text_parts)
        if header_footer_text:
            # We append header/footer text so details are searchable, but track that they were in headers/footers
            full_text = f"[HEADER/FOOTER INFO]\n{header_footer_text}\n\n[BODY]\n{full_text}"

        cleaned = clean_text(full_text)

        result = {
            "raw_text": cleaned,
            "metadata": {
                "tables_detected": tables_detected,
                "has_text_boxes": False,  # Text boxes are harder to detect inside DOCX without reading raw XML
                "is_scanned": False,      # Word docs are inherently selectable text
                "page_count": len(doc.sections),  # Section count is a rough proxy in docx
                "file_type": "docx"
            }
        }
        logger.info(f"DOCX parsed successfully. Tables found: {tables_detected}")
        return result

    except Exception as e:
        logger.error(f"Error parsing DOCX file: {e}", exc_info=True)
        raise NLPProcessingException(
            detail=f"An error occurred while parsing the DOCX document: {str(e)}",
            code="DOCX_PARSING_FAILED"
        ) from e
